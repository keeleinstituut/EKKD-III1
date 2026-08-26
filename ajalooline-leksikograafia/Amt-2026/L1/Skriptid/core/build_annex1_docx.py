#!/usr/bin/env python3
# Created: 2026-07-12 15-49-56
# Author: Madis Jürviste
# Co-Authored-By: Claude Fable 5
"""Build the thesis annex ("LISA. Väljavõte L1-andmestikust") entries as DOCX
from AMT-Master_annotated.json.

Entries only — no heading, no intro paragraphs — so the block can be pasted
straight over the entry run in the thesis WIP docx. Layout replicates the
annex as it sits in Katus-tervik-WIP (measured from Katus-20260807-WIP.docx):

- Letter page (12240 x 15840 twips), 1-inch margins, header/footer 720 twips
- two columns, 720-twip gap
- Times New Roman 11 pt (sz 22), lang et-EE, noProof on every run
- entry head: keepNext, hanging indent 220 twips, bold lemma, cross-source
  count in Cambria Math angle brackets
- DEF_et line: keepNext, flush left, italic
- et:/de: lines: hanging indent 352 twips; keepNext on all but the entry's
  last line, which instead carries 6 pt spacing after

Output is date-stamped (LISA-L1-andmestik-uuendatud-YYYYMMDD.docx) in
Katus-DRAFTS/Katus-tervik-WIP/. Grouping/label logic is shared with
build_annex1_printout.py via import.
"""

import json
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips

from build_annex1_printout import MASTER, OUT_DIR, PLACEHOLDERS, grouped_forms

DOCX = OUT_DIR / f"LISA-L1-andmestik-uuendatud-{date.today():%Y%m%d}.docx"

FONT = "Times New Roman"
MATH_FONT = "Cambria Math"
BODY_PT = Pt(11)  # sz 22 half-points


def set_two_columns(section, space_twips: int = 720) -> None:
    cols = section._sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        section._sectPr.append(cols)
    cols.set(qn("w:num"), "2")
    cols.set(qn("w:space"), str(space_twips))


def styled(par, *, hang_twips: int | None = None, after: Pt = Pt(0), keep_next=False):
    pf = par.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = after
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if keep_next:
        pf.keep_with_next = True
    if hang_twips is not None:
        pf.left_indent = Twips(hang_twips)
        pf.first_line_indent = -Twips(hang_twips)
    return par


def run(par, text: str, *, bold=False, italic=False, font=FONT):
    r = par.add_run(text)
    r.font.name = font
    rPr = r._r.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn("w:cs"), font)
    rPr.append(OxmlElement("w:noProof"))
    r.font.size = BODY_PT
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "et-EE")
    rPr.append(lang)
    return r


def main() -> None:
    data = json.load(MASTER.open())["AMT-Master"]

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_PT
    normal.paragraph_format.space_after = Pt(0)

    s = doc.sections[0]
    s.page_width, s.page_height = Twips(12240), Twips(15840)  # Letter
    for side in ("top", "bottom", "left", "right"):
        setattr(s, f"{side}_margin", Twips(1440))
    s.header_distance = Twips(720)
    s.footer_distance = Twips(720)
    set_two_columns(s, 720)

    for entry in data:
        head = styled(doc.add_paragraph(), hang_twips=220, keep_next=True)
        run(head, entry["Amt-Master-ID"], bold=True)
        run(head, " ")
        run(head, "⟨", font=MATH_FONT)
        run(head, str(entry["Cross-source count"]).strip())
        run(head, "⟩", font=MATH_FONT)

        body_pars = []
        definition = entry.get("DEF_et", "").strip()
        if definition not in PLACEHOLDERS:
            par = styled(doc.add_paragraph())  # flush left
            run(par, definition, italic=True)
            body_pars.append(par)
        for lang_code in ("et", "de"):
            parts = [
                f"{form} ({', '.join(labels)})"
                for form, labels in grouped_forms(entry, lang_code)
            ]
            if parts:
                par = styled(doc.add_paragraph(), hang_twips=352)
                run(par, f"{lang_code}: {'; '.join(parts)}")
                body_pars.append(par)
        for par in body_pars[:-1]:
            par.paragraph_format.keep_with_next = True
        if body_pars:
            body_pars[-1].paragraph_format.space_after = Pt(6)  # 120 twips
        else:
            head.paragraph_format.space_after = Pt(6)

    doc.save(DOCX)
    print(f"wrote {DOCX} ({len(data)} entries)")


if __name__ == "__main__":
    main()
